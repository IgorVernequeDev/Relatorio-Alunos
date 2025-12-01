from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import os
import re

def gerar_relatorio_docx(nome_sala, aprendizes):
    """
    Preenche o modelo modelo_relatorio.docx com os dados dos aprendizes.
    Cada aprendiz deve ser um dicionário com 'nome', 'nota' e 'observacao'.
    Mantém cada linha da tabela na mesma página e evita página em branco no final.
    """

    modelo_path = os.path.join("modelos", "modelo_relatorio.docx")
    if not os.path.exists(modelo_path):
        raise FileNotFoundError("⚠️ O arquivo 'modelo_relatorio.docx' não foi encontrado na pasta 'modelos'.")

    doc = Document(modelo_path)
    tabela = doc.tables[-1]

    while len(tabela.rows) > 1:
        tabela._tbl.remove(tabela.rows[1]._tr)

    for aprendiz in aprendizes:
        nome = str(aprendiz.get("nome") or "").strip()
        nota = str(aprendiz.get("nota") or "")
        observacao = str(aprendiz.get("observacao") or "")

        linha = tabela.add_row().cells
        linha[0].text = nome
        linha[1].text = nota
        linha[2].text = observacao

        for cell in linha:
            for paragraph in cell.paragraphs:
                pPr = paragraph._element.get_or_add_pPr()
                keepLines = OxmlElement('w:keepLines')
                pPr.append(keepLines)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    os.makedirs("relatorios", exist_ok=True)

    nome_limpo = "_".join(nome_sala.strip().split())
    nome_limpo = re.sub(r'[<>:\"/\\|?*]', '', nome_limpo)
    nome_arquivo = f"Relatório_{nome_limpo}.docx"
    caminho_saida = os.path.join("relatorios", nome_arquivo)

    doc.save(caminho_saida)
    return caminho_saida